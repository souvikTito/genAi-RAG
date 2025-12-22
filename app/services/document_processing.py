"""
Main Document processing engine for handling different file types.
Manages chunking, embeddding, preprocessing for csv/images etc.
"""

import os
import sys
import json
import unicodedata
import pandas as pd
from pathlib import Path
import fitz  # PyMuPDF library
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
from app.models.bedrock_client import invoke_claude_messages, get_bedrock_client
import re
from app.utils.util import logger, modelType, modelConfig, embeddingModel, has_embedded_newlines
import time
from datetime import datetime, timezone
import numpy as np
#Async library imports
import asyncio
from concurrent.futures import ThreadPoolExecutor
# Class imports
from typing import Optional
# Docx Imports
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph

# LangChain imports
#from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Test logger
logger.info("Starting Document Processing...")

# Simple embedding storage using Titan Embeddings for Local Run
simple_vector_store = []

executor = ThreadPoolExecutor(max_workers=5)

# Calculate max rows for CSV Truncation if needed
max_rows = modelConfig[modelType]["max_rows"]


def get_titan_embedding(text: str, bedrock_client = None, embedding_model: str = embeddingModel) -> list:
    """Get embedding using Bedrock Titan Embeddings V2"""
    
    try:
        logger.info("Invoking AWS Bedrock Embedding Module: Titan v2 (default)")
        client = bedrock_client if bedrock_client else get_bedrock_client()
  
        response = client.invoke_model(
            modelId=embedding_model,
            body=json.dumps({
                "inputText": text
            })
        )
        logger.info("Bedrock Embedding call completed")

        response_body = json.loads(response['body'].read())
        return response_body['embedding']
    except Exception as e:
        logger.error(f"Error getting Titan V2 embedding: {e}")
        return None

async def get_titan_embedding_async(text: str, bedrock_client):
    loop = asyncio.get_running_loop()
    with ThreadPoolExecutor(max_workers=5) as executor:
        return await loop.run_in_executor(executor, 
                                        lambda:  get_titan_embedding (text, bedrock_client)
                                        )

def read_pdf(file_path: str) -> str:
    """
    Extracts and preprocesses text from a PDF using PyMuPDF.
    """
    try:
        # Open the PDF document
        doc = fitz.open(file_path)
        full_text = []

        # Simple pattern for common headers/footers like page numbers
        # This will need to be customized for your specific documents
        page_pattern = re.compile(r'Page \d+ of \d+')
        
        for page in doc:
            page_text = page.get_text()

            if page_text:
                # 1. Normalize whitespace
                # Replace multiple spaces, tabs, and newlines with a single space
                page_text = re.sub(r'\s+', ' ', page_text).strip()

                # 2. Convert to lowercase
                page_text = page_text.lower()
                
                # 3. Remove headers/footers (basic example)
                page_text = re.sub(page_pattern, '', page_text)

                # 4. Remove special characters (keep alphanumerics and basic punctuation)
                page_text = re.sub(r'[^a-zA-Z0-9\s.,;:!?\'"-]', '', page_text)

                ## JSON Sanitization Optimization
                # 5.1 Remove hidden control characters (these break json.dumps)
                page_text = re.sub(r'[\x00-\x1F\x7F]', ' ', page_text)
                # 5.2 Normalize unicode (fix curly quotes, ligatures, weird forms)
                page_text = unicodedata.normalize("NFKC", page_text)
                # 5.3 Ensure valid UTF-8 (drops corrupt sequences silently)
                page_text = page_text.encode("utf-8", "ignore").decode("utf-8")

                full_text.append(page_text)
        
        doc.close()

        processed_text = ' '.join(full_text)
        
        logger.info(f"PDF processed: {len(processed_text)} characters extracted and preprocessed")
        return processed_text
    except Exception as e:
        logger.error(f"Error reading PDF with PyMuPDF: {e}")
        return ""

def read_pdf_UNUSED(file_path: str) -> str:
    """
    Extracts and preprocesses text from a PDF using pdfminer.six.
    """
    try:
        # Extract raw text from the entire PDF
        raw_text = extract_text(file_path)
        full_text = []

        # Simple pattern for common headers/footers like page numbers
        page_pattern = re.compile(r'Page \d+ of \d+')

        if raw_text:
            # 1. Normalize whitespace
            text = re.sub(r'\s+', ' ', raw_text).strip()

            # 2. Convert to lowercase
            text = text.lower()

            # 3. Remove headers/footers (basic example)
            text = re.sub(page_pattern, '', text)

            # 4. Remove special characters (keep alphanumerics and basic punctuation)
            text = re.sub(r'[^a-zA-Z0-9\s.,;:!?\'"-]', '', text)

            full_text.append(text)

        processed_text = ' '.join(full_text)
        logger.info(f"PDF processed: {len(processed_text)} characters extracted and preprocessed")
        return processed_text

    except Exception as e:
        logger.error(f"Error reading PDF with pdfminer: {e}")
        return ""

def encode_image(file_path: str) -> Optional[bytes]:
    """Read image file and return raw bytes for converse API"""
    try:
        with open(file_path, 'rb') as image_file:
            image_bytes = image_file.read()

        if not image_bytes:
            logger.error(f"Image file {file_path} was empty.")
            return None # Return None if file is unexpectedly empty
            
        logger.info(f"Image loaded as bytes: {file_path}")
        return image_bytes
 
    except Exception as e:
        logger.error(f"Error reading image: {e}")
        return None

def read_csv(file_path: str, max_rows: int = max_rows) -> str:
    """Read and preprocess CSV for RAG context"""

    try:
        logger.info("Begin CSV Read & Preprocessing...")
        try: 
            df = pd.read_csv(file_path,
                            sep = ",",
                            engine = 'python',
                            encoding = 'utf-8',
                            on_bad_lines = 'skip'
                            )
        except Exception as e:
            logger.error("Initial CSV read failed: {e}")

            # Run safety check to see if multi tab separated lines exist
            if has_embedded_newlines(file_path):
                raise ValueError (
                    "Invalid CSV format: multiline cell values detected. "
                    "Please upload a single-line CSV."
                )

            logger.info("Retrying CSV read with cp1252...")
            try:
                df = pd.read_csv(file_path,
                            sep=",",
                            engine="python",
                            encoding="cp1252",
                            errors="replace",
                            on_bad_lines="skip"
                        )
            except UnicodeDecodeError:
                logger.info("Retrying CSV read with latin1... ")
                df = pd.read_csv(file_path, 
                            encoding="latin1",
                            sep=None, engine="python", 
                            on_bad_lines="skip")
                
        
        # Debug: Log the detected columns
        logger.info(f"Detected columns by Pandas: {list(df.columns)}")
        logger.info(f"First row sample from Pandas: {df.head(1).to_dict('records')}")

        # 1. Drop completely empty rows
        df.dropna(how='all', inplace=True)

        # 2. Fill missing values with placeholder
        df.fillna("N/A", inplace=True)

        # 3. Strip whitespace and remove special characters from column names
        df.columns = [re.sub(r'[^\w\s]', '', col.strip()) for col in df.columns]

        logger.info(f"Cleaned columns by Pandas: {list(df.columns)}")
        
        # Truncate to max_rows 
        if df.shape[0] > max_rows:
            logger.info(f"CSV truncated from {df.shape[0]} to {max_rows} rows")
            df = df.head(max_rows)

        # Convert to list of dict for structured storage
        records = df.to_dict(orient="records")

        # limit rows for summary input
        limited_df = df.head()

        # Convert DataFrame to string representation    
        csv_string = limited_df.to_csv(index=False)

        # Preview Csv processed
        logger.info(f"CSV processed: {df.shape[0]} rows, {df.shape[1]} columns")

        return csv_string, records

    except Exception as e:
        logger.error(f"Error reading CSV: {e}")
        raise

def read_docx(file_path:str) -> str:
    """Extracts & preprocesses text from a DOCX file."""
    try:
        doc = Document(file_path)
        full_text = []

        # Extract paragraph text
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                full_text.append(text)

        # 2. Extract text from tables (often missed in badly formatted docs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)

        # 3. Extract text from headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for para in section.header.paragraphs:
                    text = para.text.strip()
                    if text:
                        full_text.append(text)

            # Footer
            if section.footer:
                for para in section.footer.paragraphs:
                    text = para.text.strip()
                    if text:
                        full_text.append(text)

        # 4. Extract text from text boxes and shapes (common in badly formatted docs)
        # This requires accessing the underlying XML
        try:           
            for element in doc.element.body.iter():
                # Look for text in w:txbxContent elements (text boxes)
                if element.tag.endswith('txbxContent'):
                    for text_elem in element.iter():
                        if text_elem.tag.endswith('t') and text_elem.text:
                            full_text.append(text_elem.text.strip())

        except Exception as e:
            logger.info(f"Could not extract text box content: {e}")

        # Join all text with spaces
        combined_text = ' '.join(full_text)

        # Normalise whitespace
        processed_text = re.sub(r'\s+', ' ', combined_text).strip()

        # Convert to lowercase for consistency
        processed_text = processed_text.lower() 

        # Remove unwated special characters 
        processed_text = re.sub(r'[^a-zA-Z0-9\s.,;:!?\'"-]', '', processed_text)

        logger.info(f"DOCX Processed: {len(processed_text)} characters extracted & preprocessed")
        return processed_text
    except Exception as e:
        logger.error(f"Error readong DOCX: {e}")
        raise 
    
class MultimodalPromptBuilder:
    def build(self, file_type: str, content, file_name: str, user_question: str, chat_history: list = None) -> list:
        
        messages = []

        if file_type == "image":
            # 1. Build the list of content parts for the image message
            content_parts = []
            media_type = "png" if file_name.lower().endswith(".png") else "jpeg"
            
            # Image part
            content_parts.append({ 
                "image" : { 
                    "format": media_type,
                    "source": {"bytes": content} 
                }
            })
            
            # Text part (including the initial question)
            content_parts.append({
                "text": f"Here is the image file '{file_name}'. {user_question}"
            })

            # Add the image message directly to messages
            messages.append({
                "role": "user",
                "content": content_parts
            })

        elif file_type == "csv":
            system_context = (
                f"{user_question}.\n"
                f"CSV Context:\n{content}\n"
            )
        else:
            raise ValueError(f"Unsupported file type for multimodal prompt: {file_type}")

        # Add previous chat history if present
        if chat_history:
            for msg in chat_history:
                role = "user" if msg.get("role") == "human" else "assistant"
                messages.append({
                    "role": role,
                    "content": [{"text": msg.get("content", "")}]
                })

        # Add current user question (common)
        if file_type == "csv":
            # Prepend system context to user question for CSV
            final_user_message = system_context + user_question
        else:
            final_user_message = user_question

        # Add current user question
        messages.append({
            "role": "user",
            "content": [{"text": final_user_message}]
        })

        return messages
        
async def store_embeddings_async(content: str, file_type: str, source: str, bedrock_client,
                                 chunk_size=3000, chunk_overlap=500,
                                 separators=None):
    separators = separators or ["\n\n", "\n", ". ", " ", ""]
    
    # Create a fresh event loop
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)

    # Split content into chunks inside the async function
    logger.info("Starting chunking process using RecursiveCharacterTextSplitter")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=separators
    )
    chunks = text_splitter.split_text(content)

    # Time monitoring for Embedding
    embedding_start_time = time.time()
    logger.info("Starting embedding process for chunks")

    embeddings = []
    # Run embeddings in parallel
    embeddings = await asyncio.gather(*[get_titan_embedding_async(c, bedrock_client) for c in chunks])

    embedding_end_time = time.time()
    embedding_duration = embedding_end_time - embedding_start_time
    logger.info(f"Embedding completed in {embedding_duration:.2f} seconds")

    for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        if embedding is None:
            logger.warning(f"Failed to create embedding for chunk {i+1}")
            continue

        embedding_data = {
            'content': chunk,
            'embedding': embedding,
            'file_type': file_type,
            'source': f"{source} (chunk {i+1}/{len(chunks)})",
            'id': len(simple_vector_store)
        }
        simple_vector_store.append(embedding_data)

    logger.info(f"Successfully stored {len(chunks)} Titan V2 embeddings for {file_type}: {source}")
    return simple_vector_store

def search_content(query: str, top_k: int = 3) -> str:
    """Semantic search using Titan embeddings"""
    # Time monitoring for Retrieval
    retrieval_start_time = time.time()
    logger.info("Starting retrieval process")
    
    if not simple_vector_store:
        logger.warning("No content stored yet")
        return "No content stored yet."
    
    # Get query embedding using Titan
    query_embedding = get_titan_embedding(query)
    if query_embedding is None:
        logger.error("Error creating query embedding")
        return "Error creating query embedding."
    
    # Calculate similarities
    similarities = []
    for item in simple_vector_store:
        similarity = cosine_similarity([query_embedding], [item['embedding']])[0][0]
        similarities.append({
            'content': item['content'],
            'similarity': similarity,
            'source': item['source'],
            'file_type': item['file_type']
        })
    
    # Sort by similarity and get top results
    similarities.sort(key=lambda x: x['similarity'], reverse=True)
    
    retrieval_end_time = time.time()
    retrieval_duration = retrieval_end_time - retrieval_start_time
    logger.info(f"Retrieval completed in {retrieval_duration:.2f} seconds")
    
    '''
    # Print the highest scoring chunk
    if similarities:
        best_match = similarities[0]
        logger.info(f"\n=== HIGHEST SIMILARITY CHUNK ===")
        logger.info(f"Source: {best_match['source']}")
        logger.info(f"Similarity: {best_match['similarity']:.3f}")
        logger.info(f"Content ({len(best_match['content'])} chars):")
        logger.info("-" * 80)
        logger.info(best_match['content'])
        logger.info("-" * 80)
    '''

    relevant_content = []
    for result in similarities[:top_k]:
        #logger.info(f"Match: {result['source']} (similarity: {result['similarity']:.3f})")
        relevant_content.append(f"From {result['source']} ({result['file_type']}):\n{result['content'][:2000]}...")
    
    return "\n\n---\n\n".join(relevant_content)

def create_rag_prompt_UNUSED(user_question: str, context: str) -> str:
    """Create a simple RAG prompt template"""
    prompt = f"""Based on the following context information, please answer the user's question.

Context:
{context}

Question: {user_question}

Please provide a helpful answer based on the context provided. If the context doesn't contain relevant information, please say so."""
    
    return prompt

def create_rag_prompt(user_question: str, context: str, chat_history: list = None) -> list:
    """
    Creates a production-ready RAG prompt,
    with an optional chat history.
    """
    # System message with context
    messages = [
        {
            "role": "system",
            "content": [{"type": "text", "text": (
                "You are a helpful AI assistant. Use the following pieces of context to answer the user's question. "
                "If you don't know the answer, just say that you don't know. Keep your answer concise and do not make up an answer.\n\n"
                f"Context: {context}"
            )}]
        }
    ]

    # Add previous chat history if present
    if chat_history:
        for msg in chat_history:
            role = "user" if msg.get("role") == "human" else "assistant"
            messages.append({
                "role": role,
                "content": [{"type": "text", "text": msg.get("content", "")}]
            })

    # Add current user question
    messages.append({
        "role": "user",
        "content": [{"type": "text", "text": user_question}]
    })

    return messages

def process_file(file_path: str, bedrock_client, docId: str = "test123456",
                 user_question: str = "Provide a detailed description of the document or a sample content of the document in plain text without bullet points, numbered lists, or special formatting. Be direct and concise without introductory phrases.",
                 model_type: str = modelType, embedding_model: str = embeddingModel):
    
    """Process different file types"""
    logger.info(f"Starting Process File Calls for {docId}")
    start_time = datetime.now(timezone.utc)

    file_path = Path(file_path)
    logger.info(f"Fetching file from {file_path}")

    if not file_path.exists():
        logger.error(f"File not found: {file_path}")
        return []
    
    
    if file_path.suffix.lower() == '.pdf':
        content = read_pdf(str(file_path))
        logger.info("File loaded from Path: ", file_path)
        if not content:
            logger.warning(f"No content extracted from PDF: {file_path}")
            return None

        else:
            vectors = asyncio.run(store_embeddings_async(content, 'pdf', str(file_path), bedrock_client))
            #return vectors # to directly return embeddings vector

            #Generate summary
            summary_content = content[:10000] if len(content) > 10000 else content
            summary_prompt = [{
                "role": "user",
                "content": [{"text": f"{user_question}\n{summary_content}"}]
            }]

            try:
                summary = invoke_claude_messages(summary_prompt, max_tokens=500)
                logger.info(f"Document Summary: {summary}")
            except Exception as e:
                logger.error(f"Failed to generate summary: {e}")
                summary = "Summary generation failed"

            end_time = datetime.now(timezone.utc)
            response_time = end_time - start_time
            logger.info(f"Document Processing Time: {round(response_time.total_seconds(), 2)}")

            # Return both vectors and summary
            return {
                "docId": docId,
                "file_type": file_path.suffix.lower(),
                "source": str(file_path),
                "summary": summary,
                "chunks": vectors,
                "embeddings": embedding_model, # [chunk["embedding"] for chunk in vectors]     # to return embeddings list          
                "response_time": response_time
            }

    elif file_path.suffix.lower() == '.docx':
        content = read_docx(str(file_path))
        logger.info("File loaded from Path: ", file_path)
        
        if not content:
            logger.warning(f"No content extracted from PDF: {file_path}")
            return None
        else:
            vectors = asyncio.run(store_embeddings_async(content, 'docx', str(file_path), bedrock_client))
            
            #Generate summary
            summary_content = content[:10000] if len(content) > 10000 else content
            summary_prompt = [{
                "role": "user",
                "content": [{"text": f"{user_question}\n{summary_content}"}]
            }]
            
            try:
                summary = invoke_claude_messages(summary_prompt, max_tokens=500)
                logger.info(f"Document Summary: {summary}")
            except Exception as e:
                logger.error(f"Failed to generate summary: {e}")
                summary = "Summary generation failed"

            end_time = datetime.now(timezone.utc)
            response_time = end_time - start_time
            logger.info(f"Document Processing Time: {round(response_time.total_seconds(), 2)}")

            # Return both vectors and summary
            return {
                "docId": docId,
                "file_type": file_path.suffix.lower(),
                "source": str(file_path),
                "summary": summary,
                "chunks": vectors,
                "embeddings":embedding_model, # [chunk["embedding"] for chunk in vectors]     # to return embeddings list          
                "response_time": response_time
            }

    elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png']:
        #if model_type.lower() == "haiku":
        #    logger.error(
        #        f"Skipping image file {file_path.name}. "
        #        f"Claude {model_type} (Model ID: {modelId}) does not support multimodal (image) input. "
        #        "Please configure a Sonnet or supported model to process images."
        #    )
        #    return None # Exit gracefully for unsupported model/file combination
        # --------------------------------
        
        encoded_image = encode_image(str(file_path))
        if encoded_image:
            builder = MultimodalPromptBuilder()
            image_prompt = builder.build(
                file_type="image",
                content=encoded_image, 
                file_name=file_path.name,
                user_question=user_question
            )
            
            # Generate image summary
            try:
                summary = invoke_claude_messages(image_prompt, max_tokens=500)
                logger.info(f"Document Summary: {summary}")
            except Exception as e:
                logger.error(f"Failed to generate summary: {e}")
                summary = "Summary generation failed"

            # For Local run -append im memory store
            simple_vector_store.append({
                "file_type": "image",
                "source": str(file_path),
                "prompt": image_prompt,
                "image_data": encoded_image
            })

            end_time = datetime.now(timezone.utc)
            response_time = end_time - start_time
            logger.info(f"Document Processing Time: {round(response_time.total_seconds(), 2)}")    

            return {
                "docId": docId,
                "file_type": "image",
                "source": str(file_path),
                "summary": summary,
                "image_data": encoded_image,
                "response_time": response_time
            }

    elif file_path.suffix.lower() == '.csv':
        content, records = read_csv(str(file_path))
        if content:
            builder = MultimodalPromptBuilder()
            csv_prompt = builder.build(
                file_type="csv",
                content=content,
                file_name=file_path.name,
                user_question=user_question
            )

            # Generate CSV summary
            try:
                summary = invoke_claude_messages(csv_prompt, max_tokens=500)
                logger.info(f"Document Summary: {summary}")
            except Exception as e:
                logger.error(f"Failed to generate summary: {e}")
                summary = "Summary generation failed"
            
            # For Local run -append im memory store
            simple_vector_store.append({
                "file_type": "csv",
                "source": str(file_path),
                "prompt": csv_prompt
            })


            end_time = datetime.now(timezone.utc)
            response_time = end_time - start_time
            logger.info(f"Document Processing Time: {round(response_time.total_seconds(), 2)}")

            return {
                "docId": docId,
                "file_type": "csv",
                "source": str(file_path),
                "summary": summary,
                "content": records,
                "response_time": round(response_time.total_seconds(), 2)
            }        
    else:
        logger.warning(f"Unsupported file type: {file_path.suffix}")
        return None
    
    return None

def rag_query(question: str) -> str:
    """Alternative RAG query message structure"""
    # Search for relevant content
    context = search_content(question)
    
    # Create messages structure
    messages = create_rag_prompt(question, context)
    
    # Convert to the format expected by bedrock.converse()
    bedrock_messages = []
    for msg in messages:
        if msg["role"] == "system":
            # Bedrock converse doesn't use system role, so we'll prepend to first user message
            continue
        elif msg["role"] == "user":
            bedrock_messages.append({
                "role": "user",
                "content": [{"text": msg["content"][0]["text"]}]
            })
        elif msg["role"] == "assistant":
            bedrock_messages.append({
                "role": "assistant", 
                "content": [{"text": msg["content"][0]["text"]}]
            })
    
    # Add system context to the first user message
    if messages and messages[0]["role"] == "system":
        system_content = messages[0]["content"][0]["text"]
        if bedrock_messages and bedrock_messages[0]["role"] == "user":
            bedrock_messages[0]["content"][0]["text"] = system_content + "\n\n" + bedrock_messages[0]["content"][0]["text"]
    
    # Debug print
    #logger.debug(f"\n=== BEDROCK MESSAGES ===")
    #logger.debug(json.dumps(bedrock_messages, indent=2))
    #logger.debug("=" * 40)
    
    # Use the modified invoke function
    return invoke_claude_messages(bedrock_messages)

# Test function
def test_rag_system(question: str = "Provide a detailed description of the document in plain text without bullet points, numbered lists, or special formatting. Be direct and concise without introductory phrases.",
                    ):
    """Simple test of the RAG system"""
    logger.info("=== Bedrock RAG POC Test ===")
    
    # Example usage
    test_files = [
        r"C:\Users\600002608\medpro-data-ai-gpt\genai\app\testdocs\test_medpro.docx", 
        #"sample.jpg", 
        #"sample.csv" 
    ]
    
    # Process files
    for file_path in test_files:
        if os.path.exists(file_path):
            logger.info(f"Processing: {file_path}")
            try:
                was_successful = process_file(file_path, user_question=question, model_type=modelType)

                if was_successful:  
                    logger.info(f"Successfully processed: {file_path}")
                else:
                    logger.warning(f"File processing skipped or failed: {file_path}. Skipping query for this file type.")
                    continue
            except Exception as e:
                logger.error(f"FATAL File Processing Error during {file_path}: {e}")
                # Log the full traceback for maximum debug info
                logger.error("Traceback follows:", exc_info=True) 
                sys.exit(1)
        else:
            logger.warning(f"File not found (skipping): {file_path}")
    


    # Test single specific query
    if simple_vector_store:
        # Show what's stored
        logger.info(f"Stored {len(simple_vector_store)} items in vector store")
        
        stored_item = simple_vector_store[0]    
        file_type = stored_item.get("file_type")
        logger.info("=== Testing RAG Query ===")
        logger.info(f"Q: {question}")

        if file_type in ["pdf","docx"]:
            answer = rag_query(question)
            logger.info(f"A: {answer}")
        
        elif file_type in ["csv", "image"]:
            prompt_with_context = simple_vector_store[0]["prompt"]
            # Create and append the final user query message
            final_user_query_message = {
            "role": "user",
            "content": [{"text": question}] 
                }
            
            final_messages = prompt_with_context + [final_user_query_message]

            # Use the complete message list for invocation
            response = invoke_claude_messages(final_messages, max_tokens=200)

            logger.info(f"A: {response}")

    else:
        logger.info("No files processed - add some test files to try the RAG system")


if __name__ == "__main__":
    # Run sample RAG operation
    #test_rag_system("What are the trends in the housing values based on population?") #for csv
    #test_rag_system("What is available in terms of antibiotics in the document?") # for McDonalds.pdf
    test_rag_system()

    
